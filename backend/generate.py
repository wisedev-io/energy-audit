"""
Document generation: Excel copy, Word report, Energy Passport, charts.
"""

import copy
import os
import random
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

os.environ.setdefault(
    'MPLCONFIGDIR',
    str(Path(tempfile.gettempdir()) / 'energy-audit-mpl'),
)

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from calculate import (
    MONTHS, calc_all, calc_areas, calc_ariston, calc_efficiency,
    fes_values, gelio_values,
)
from storage import BinaryUpload, get_storage, guess_mimetype

BASE_DIR  = Path(__file__).parent


# ── Charts ─────────────────────────────────────────────────────────────────

def _fmt_mln(x, _):
    return f"{x/1e6:.1f}"

def generate_charts(energy, out_dir):
    """Generate 4 charts: 3 pie charts (one per year) + 1 bar chart.
    Returns list of PNG paths."""
    paths = []
    years = [2023, 2024, 2025]
    colors = ['#4472C4', '#ED7D31', '#A9D18E']
    labels = ['Elektr', 'Gaz', 'Boshqa']

    # --- 3 Pie charts ---
    for year in years:
        s = energy[year]['summary']
        sizes = [s['total_elec'], s['total_gkwh'], s['total_okwh']]
        if sum(sizes) == 0:
            sizes = [1, 1, 1]
        fig, ax = plt.subplots(figsize=(4.5, 4))
        wedges, texts, autotexts = ax.pie(
            sizes, labels=labels, colors=colors,
            autopct='%1.1f%%', startangle=140,
            pctdistance=0.65, labeldistance=1.18,
            textprops={'fontsize': 9},
        )
        for at in autotexts:
            at.set_fontsize(8)
            at.set_fontweight('bold')
        ax.set_title(f"{year}-yil energiya ulushi", fontsize=10, fontweight='bold', pad=8)
        plt.tight_layout(pad=0.5)
        p = out_dir / f'chart_pie_{year}.png'
        fig.savefig(p, dpi=150, bbox_inches='tight')
        plt.close(fig)
        paths.append(p)

    # --- Bar chart: annual total cost ---
    totals = [energy[y]['summary']['total_cost'] for y in years]
    fig, ax = plt.subplots(figsize=(5, 3.5))
    bars = ax.bar([str(y) for y in years], totals, color=colors[:3], width=0.5)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(_fmt_mln))
    ax.set_ylabel("mln. so'm", fontsize=9)
    ax.set_title("Yillik energiya xarajatlari (mln. so'm)", fontsize=10, fontweight='bold')
    for bar, val in zip(bars, totals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(totals)*0.01,
                f"{val/1e6:.2f}", ha='center', va='bottom', fontsize=8)
    plt.tight_layout()
    p = out_dir / 'chart_bar_cost.png'
    fig.savefig(p, dpi=150, bbox_inches='tight')
    plt.close(fig)
    paths.append(p)

    return paths


# ── Excel output ───────────────────────────────────────────────────────────

def write_excel(data, energy, areas, out_path):
    src = BASE_DIR / 'template.xlsx'
    dst = Path(out_path)
    shutil.copy2(src, dst)

    wb = openpyxl.load_workbook(dst)
    ws = wb['Main']

    # ── Raw monthly inputs (rows 2-13, cols A-J) ──────────────────────────
    # A=gas2023, B=gas2024, C=gas2025, D=elec2023, E=elec2024, F=elec2025
    # G=elec2026(empty), H=other2023, I=other2024, J=other2025
    keys = ['gas_2023','gas_2024','gas_2025','elec_2023','elec_2024','elec_2025',
            'elec_2026','other_2023','other_2024','other_2025']
    for col_idx, key in enumerate(keys):
        vals = data.get(key, [0]*12)
        for row_idx in range(12):
            ws.cell(row=row_idx + 2, column=col_idx + 1,
                    value=float(vals[row_idx] or 0))

    # ── Per-year display sections ──────────────────────────────────────────
    # base_row = first month row; summary_row = row in K-O comparison table
    YEAR_LAYOUT = [(2023, 17, 76), (2024, 36, 77), (2025, 55, 78)]

    for year, base_row, cmp_row in YEAR_LAYOUT:
        rows = energy[year]['rows']
        s    = energy[year]['summary']

        # Monthly rows (base_row .. base_row+11): cols B-G
        for i, r in enumerate(rows):
            ws.cell(row=base_row + i, column=2, value=int(round(r['elec_kwh'])))
            ws.cell(row=base_row + i, column=3, value=int(round(r['gas_m3'])))
            ws.cell(row=base_row + i, column=4, value=round(r['gas_kwh'],   1))
            ws.cell(row=base_row + i, column=5, value=int(round(r['other_kg'])))
            ws.cell(row=base_row + i, column=6, value=round(r['other_kwh'], 1))
            ws.cell(row=base_row + i, column=7, value=round(r['total_kwh'], 1))

        # JAMI row (base_row+12)
        jr = base_row + 12
        ws.cell(row=jr, column=2, value=int(s['total_elec']))
        ws.cell(row=jr, column=3, value=int(s['total_gas']))
        ws.cell(row=jr, column=4, value=round(s['total_gkwh'], 1))
        ws.cell(row=jr, column=5, value=int(s['total_okg']))
        ws.cell(row=jr, column=6, value=round(s['total_okwh'], 1))
        ws.cell(row=jr, column=7, value=round(s['total_kwh'],  1))

        # "JAMI yillik energiya sarfi" row (base_row+13): col C = total kWh
        ws.cell(row=base_row + 13, column=3, value=round(s['total_kwh'], 1))

        # "Ulushi" row (base_row+14): fractions — cell format is "0%" so 0.65 → 65%
        ws.cell(row=base_row + 14, column=2, value=round(s['pct_gas']   / 100, 4))
        ws.cell(row=base_row + 14, column=4, value=round(s['pct_elec']  / 100, 4))
        ws.cell(row=base_row + 14, column=6, value=round(s['pct_other'] / 100, 4))

        # "Yillik umumiy xarajatlar" row (base_row+15): col C = cost in mln so'm
        ws.cell(row=base_row + 15, column=3, value=round(s['total_cost'] / 1_000_000, 4))

        # Comparison/summary table row (cmp_row): cols K(11)–R(18)
        ws.cell(row=cmp_row, column=11, value=round(s['total_kwh'],  1))   # K total kWh
        ws.cell(row=cmp_row, column=12, value=round(s['pct_elec']  / 100, 4))  # L elec %
        ws.cell(row=cmp_row, column=13, value=round(s['pct_gas']   / 100, 4))  # M gas %
        ws.cell(row=cmp_row, column=14, value=round(s['pct_other'] / 100, 4))  # N other %
        ws.cell(row=cmp_row, column=15, value=int(s['total_cost']))             # O cost som
        ws.cell(row=cmp_row, column=16, value=round(s['total_elec'],  1))  # P elec kWh
        ws.cell(row=cmp_row, column=17, value=round(s['total_gkwh'],  1))  # Q gas kWh
        ws.cell(row=cmp_row, column=18, value=round(s['total_okwh'],  1))  # R other kWh

    # Annual cost summary (rows 81-83 col E): mln so'm per year
    for row_e, year in [(81, 2023), (82, 2024), (83, 2025)]:
        ws.cell(row=row_e, column=5,
                value=round(energy[year]['summary']['total_cost'] / 1_000_000, 4))

    # Force Excel to recalculate remaining formula cells on open
    wb.calculation.calcMode = 'auto'

    # ── Area_calc sheet ───────────────────────────────────────────────────
    if 'Area_calc' in wb.sheetnames:
        wa = wb['Area_calc']
        for i in range(1, 6):
            wa[f'A{i+2}'] = float(data.get(f'floor_l{i}') or 0)
            wa[f'B{i+2}'] = float(data.get(f'floor_w{i}') or 0)
        for i in range(1, 6):
            wa[f'E{i+2}'] = float(data.get(f'win_w{i}') or 0)
            wa[f'F{i+2}'] = float(data.get(f'win_h{i}') or 0)
            wa[f'G{i+2}'] = float(data.get(f'win_n{i}') or 0)
        for i in range(1, 4):
            wa[f'E{i+18}'] = float(data.get(f'door_w{i}') or 0)
            wa[f'F{i+18}'] = float(data.get(f'door_h{i}') or 0)
            wa[f'G{i+18}'] = float(data.get(f'door_n{i}') or 0)
        wa['B36'] = areas['roof_area']

    wb.save(dst)
    return dst


# ── Word helpers ───────────────────────────────────────────────────────────

def _replace_in_para(para, variables):
    """Replace {{key}} placeholders, skipping runs that contain images."""
    full = ''.join(r.text for r in para.runs)
    if '{{' not in full:
        return
    for key, val in variables.items():
        full = full.replace('{{' + key + '}}', str(val))
    # Only touch runs that have no drawing/image content
    text_runs = [r for r in para.runs
                 if '<w:drawing' not in r._element.xml and '<v:' not in r._element.xml]
    if not text_runs:
        return
    text_runs[0].text = full
    for r in text_runs[1:]:
        r.text = ''


def _replace_in_cell(cell, variables):
    for para in cell.paragraphs:
        _replace_in_para(para, variables)


def _replace_in_header_footer(hf, variables):
    """Replace placeholders in a header or footer object."""
    for para in hf.paragraphs:
        _replace_in_para(para, variables)
    for table in hf.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, variables)


def _replace_all(doc, variables):
    # Body paragraphs and tables
    for para in doc.paragraphs:
        _replace_in_para(para, variables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, variables)
    # Headers and footers in every section
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            try:
                _replace_in_header_footer(hf, variables)
            except Exception:
                pass


def _set_cell_text(cell, text, bold=False, font_size=None, align=None, font_name=None):
    for para in cell.paragraphs:
        para.clear()
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if align:
        para.alignment = align


def _cell_content_width_inches(table, col_idx=0):
    """Return usable content width in inches for a table column (minus cell margins)."""
    from docx.oxml.ns import qn
    TWIPS_PER_INCH = 1440
    CELL_MARGIN_TWIPS = 108  # Word default per side (108 left + 108 right = 216 total)
    try:
        tbl_grid = table._tbl.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            grid_cols = tbl_grid.findall(qn('w:gridCol'))
            if 0 <= col_idx < len(grid_cols):
                w_twips = int(grid_cols[col_idx].get(qn('w:w'), 0))
                if w_twips > 0:
                    return max(0.5, round((w_twips - 2 * CELL_MARGIN_TWIPS) / TWIPS_PER_INCH, 3))
    except Exception:
        pass
    return None


def _cell_actual_width_inches(cell):
    """Return cell width in inches from its own XML, correctly handles merged columns."""
    from docx.oxml.ns import qn
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None:
            w = tcW.get(qn('w:w'))
            typ = tcW.get(qn('w:type'))
            if w and typ == 'dxa' and int(w) > 0:
                return int(w) / 1440.0
    return None

def _col_width_from_tblGrid(table, col_idx):
    """Get column width in inches from table grid definition."""
    from docx.oxml.ns import qn
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        cols = tblGrid.findall(qn('w:gridCol'))
        if col_idx < len(cols):
            w = cols[col_idx].get(qn('w:w'))
            if w:
                return int(w) / 1440.0
    return None

def _insert_image_in_cell(cell, img_path, table=None, col_idx=0, width_in=None):
    """Insert image scaled to fill cell width exactly."""
    import traceback

    img_path_obj = Path(img_path)
    if not img_path_obj.exists():
        print(f"⚠ _insert_image_in_cell: file does not exist: {img_path}")
        return

    file_size = img_path_obj.stat().st_size
    if file_size == 0:
        print(f"⚠ _insert_image_in_cell: file is empty (0 bytes): {img_path}")
        return

    try:
        # Priority: explicit width → cell XML → tblGrid → fallback
        w = width_in
        if w is None:
            w = _cell_actual_width_inches(cell)
        if w is None and table is not None:
            w = _col_width_from_tblGrid(table, col_idx)
        if w is None and table is not None:
            w = _cell_content_width_inches(table, col_idx)
        if w is None:
            w = 3.0
        # subtract cell left+right margins (default 0.08 in each side)
        final_w = max(0.5, w - 0.16)
        print(f"  inserting {img_path_obj.name} ({file_size} bytes) at {final_w:.2f}in (col {col_idx})")

        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run()
        run.add_picture(str(img_path), width=Inches(final_w))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"⚠ _insert_image_in_cell error for {img_path}: {e}")
        traceback.print_exc()


def _insert_image_after_para(doc, para_idx, img_path, caption, width_in=5.5):
    """Insert image paragraph after a given paragraph index."""
    # Find the paragraph in the document body
    body = doc.element.body
    paras = doc.paragraphs
    if para_idx >= len(paras):
        return
    ref_para = paras[para_idx]._element

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    # Create new paragraph for image
    new_para = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_para.append(new_r)

    # Insert after reference paragraph
    ref_para.addnext(new_para)

    # Now use python-docx to add the picture
    # We need to find the new paragraph object
    for i, p in enumerate(doc.paragraphs):
        if p._element is new_para:
            run = p.add_run()
            try:
                run.add_picture(str(img_path), width=Inches(width_in))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                run.text = f"[Chart: {Path(img_path).name}]"
            break


# ── Fill energy tables in Word ─────────────────────────────────────────────

def _fill_energy_table(table, rows, summary, year):
    """Fill a monthly energy Word table from calculated rows."""
    # Detect column count to handle 2023 (8 cols) vs 2024/2025 (7 cols)
    def unique_cells(row):
        seen, result = set(), []
        for c in row.cells:
            if id(c._tc) not in seen:
                seen.add(id(c._tc))
                result.append(c)
        return result

    def set_c(cells, idx, val, bold=False):
        if idx < len(cells):
            _set_cell_text(cells[idx], val, bold=bold, font_size=11, font_name='Cambria')

    for i, r in enumerate(rows):
        row_idx = i + 1
        if row_idx >= len(table.rows):
            break
        cells = unique_cells(table.rows[row_idx])
        set_c(cells, 0, r['month'])
        set_c(cells, 1, int(r['elec_kwh']))
        set_c(cells, 2, int(r['gas_m3']))
        set_c(cells, 3, round(r['gas_kwh'], 1))
        set_c(cells, 4, int(r['other_kg']))
        set_c(cells, 5, round(r['other_kwh'], 1))
        set_c(cells, 6, round(r['total_kwh'], 1))

    if 13 < len(table.rows):
        cells = unique_cells(table.rows[13])
        set_c(cells, 0, 'JAMI',                          bold=True)
        set_c(cells, 1, int(summary['total_elec']),       bold=True)
        set_c(cells, 2, int(summary['total_gas']),        bold=True)
        set_c(cells, 3, round(summary['total_gkwh'], 1),  bold=True)
        set_c(cells, 4, int(summary['total_okg']),        bold=True)
        set_c(cells, 5, round(summary['total_okwh'], 1),  bold=True)
        set_c(cells, 6, round(summary['total_kwh'],  1),  bold=True)

    # Summary rows 14-16: find empty cells between label and unit text
    def _fill_between(row_idx, before_label, after_label, value, bold=False):
        """Fill the first empty unique cell between two labelled cells."""
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        seen, past_before = set(), False
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            txt = cell.text.strip()
            if before_label and before_label in txt:
                past_before = True
                continue
            if past_before or not before_label:
                if after_label and after_label in txt:
                    break
                if not txt or txt == '\xa0':
                    _set_cell_text(cell, str(value), bold=bold, font_size=11, font_name='Cambria')
                    return

    _fill_between(14, 'JAMI yillik',  'kW',     int(summary['total_kwh']), bold=True)
    _fill_between(15, 'Ulushi',       'gaz',    f"{summary['pct_gas']}%",   bold=True)
    _fill_between(15, 'gaz',          'elektr', f"{summary['pct_elec']}%",  bold=True)
    _fill_between(15, 'elektr',       'boshqa', f"{summary['pct_other']}%", bold=True)
    _fill_between(16, 'Yillik',       'mln',    f"{summary['total_mln']:.4f}", bold=True)


def _fill_comparison_table(table, energy):
    """Fill Table 16: rows 1-3 comparison, rows 7-9 pie data 2023/2024, rows 13-15 pie data 2025/costs."""

    def sc(row_idx, col_idx, val, bold=False):
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        seen, unique = set(), []
        for c in row.cells:
            if id(c._tc) not in seen:
                seen.add(id(c._tc))
                unique.append(c)
        if col_idx < len(unique):
            _set_cell_text(unique[col_idx], val, bold=bold, font_size=11, font_name='Cambria')

    years = [2023, 2024, 2025]
    kwh_keys = ['total_elec', 'total_gkwh', 'total_okwh']
    pct_keys = ['pct_elec', 'pct_gas', 'pct_other']

    # Rows 1-3: main comparison table
    for i, year in enumerate(years):
        s = energy[year]['summary']
        ri = i + 1
        sc(ri, 0, f"{year}-yil")
        sc(ri, 1, int(s['total_kwh']))
        sc(ri, 2, f"{s['pct_elec']}%")
        sc(ri, 3, f"{s['pct_gas']}%")
        sc(ri, 4, f"{s['pct_other']}%")
        sc(ri, 5, f"{s['total_mln']:.4f}")
        sc(ri, 6, int(s['total_elec']))
        sc(ri, 7, int(s['total_gkwh']))

    # Rows 7-9: 2023 pie data cols 1-2, 2024 pie data cols 5-6
    for j in range(3):
        ri = 7 + j
        s23 = energy[2023]['summary']
        s24 = energy[2024]['summary']
        sc(ri, 1, int(s23[kwh_keys[j]]))
        sc(ri, 2, f"{s23[pct_keys[j]]}%")
        sc(ri, 5, int(s24[kwh_keys[j]]))
        sc(ri, 6, f"{s24[pct_keys[j]]}%")

    # Rows 13-15: 2025 pie data cols 1-2, jami xarajat col 5
    for j in range(3):
        ri = 13 + j
        s25 = energy[2025]['summary']
        sc(ri, 1, int(s25[kwh_keys[j]]))
        sc(ri, 2, f"{s25[pct_keys[j]]}%")

    cost_years = [(13, 2023), (14, 2024), (15, 2025)]
    for ri, year in cost_years:
        s = energy[year]['summary']
        sc(ri, 5, f"{s['total_mln']:.4f}")


# ── Multi-photo section definitions ───────────────────────────────────────
# sec_id → (table_idx, [placeholder_row_idxs_to_remove], caption_row_idx)
# photo_row_idxs: all empty placeholder rows that will be removed before inserting real photos
# caption_row_idx: the merged caption row immediately after the photo area
PHOTO_SECTION_MAP = {
    # NOTE: Indices here are tied to `template.docx` table order.
    1:  (3,  [0],  1),  # 1-rasm: exterior view; table 3 row 0=photo, row 1=caption
    2:  (3,  [2],  3),  # 2-rasm: doors & windows; table 3 row 2=photo, row 3=caption
    3:  (3,  [4],  5),  # 3-rasm: floor plan; table 3 row 4=photo, row 5=caption
    4:  (5,  [0],        1),  # 4-rasm: heating system (up to 2; table has 2 cols)
    5:  (5,  [2],        3),  # 5-rasm: AC / cooling (up to 1; table has 2 cols)
    6:  (7,  [0, 1, 2],  3),  # 6-rasm: appliances (up to 6; table has 2 cols)
    7:  (8,  [0, 1, 2],  3),  # 7-rasm: energy bill screenshots (up to 6; table has 3 cols)
    8:  (19, [0],  1),  # 9-rasm: temperature & humidity; table 19 row0=photo row1=caption
    9:  (21, [0],  1),  # 10-rasm: lux measurement; table 21 row0=photo row1=caption
    # 10 = thermal camera → Table 25 special handling (columns 3-4)
}

DEFAULT_PHOTO_WIDTHS = {
    1: 3.0, 2: 3.0, 3: 3.0, 4: 3.0, 5: 3.0,
    6: 3.0, 7: 3.0, 8: 3.0, 9: 3.0, 10: 3.0,
}


def _make_tr_copy(template_tr):
    """Return a deep copy of a TR element with all cell contents cleared."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    new_tr = copy.deepcopy(template_tr)
    for tc in new_tr.findall(qn('w:tc')):
        for p in list(tc.findall(qn('w:p'))):
            tc.remove(p)
        tc.append(OxmlElement('w:p'))
    return new_tr


def _convert_image_if_needed(img_path):
    """Convert HEIC/HEIF to JPG. Returns the final image path."""
    if img_path.suffix.lower() not in ['.heic', '.heif']:
        return img_path

    jpg_path = img_path.with_suffix('.jpg')
    try:
        import pillow_heif
        from PIL import Image
        pillow_heif.register_heif_opener()
        img = Image.open(str(img_path))
        img.convert('RGB').save(str(jpg_path), 'JPEG', quality=90)
        img_path.unlink()
        return jpg_path
    except Exception as e:
        try:
            subprocess.run(
                ['sips', '-s', 'format', 'jpeg', str(img_path), '--out', str(jpg_path)],
                check=True,
                capture_output=True,
            )
            img_path.unlink()
            return jpg_path
        except Exception as sips_err:
            print(f"⚠ HEIC konvertatsiya xatosi {img_path.name}: {e}; sips fallback: {sips_err}")
            return img_path


def _normalize_to_jpeg(img_path: Path) -> Path:
    """Re-save any image as JPEG via PIL, catching format-mismatch issues.

    iPhone photos often arrive with .jpg extension but contain HEIC/WebP data;
    PIL detects format by content, so this always produces a valid JPEG that
    python-docx can embed without error.
    """
    import traceback
    from PIL import Image as PilImage

    # Register HEIC/HEIF opener if available (handles iPhone photos)
    try:
        import pillow_heif
        pillow_heif.register_heif_opener()
    except ImportError:
        pass

    jpg_path = img_path.with_suffix('.jpg')
    try:
        with PilImage.open(str(img_path)) as im:
            fmt = im.format
            print(f"  _normalize_to_jpeg: {img_path.name} detected as {fmt}, size={im.size}")
            im.convert('RGB').save(str(jpg_path), 'JPEG', quality=90)
        if img_path != jpg_path and img_path.exists():
            img_path.unlink()
        print(f"  _normalize_to_jpeg: saved → {jpg_path.name} ({jpg_path.stat().st_size} bytes)")
        return jpg_path
    except Exception as e:
        print(f"⚠ _normalize_to_jpeg failed for {img_path.name}: {e}")
        traceback.print_exc()
        return img_path


def _fill_photo_table_section(doc, sec_id, photos, width_in=None):
    """Append uploaded photos after the template examples for one section."""
    from docx.oxml.ns import qn
    from docx.table import _Cell

    if not photos or sec_id not in PHOTO_SECTION_MAP:
        return

    t_idx, placeholder_row_idxs, caption_row_idx = PHOTO_SECTION_MAP[sec_id]
    if t_idx >= len(doc.tables):
        return
    table = doc.tables[t_idx]
    n_cols = len(table.columns)

    if caption_row_idx >= len(table.rows):
        print(f"\u26a0 sec {sec_id}: caption_row_idx={caption_row_idx} >= rows={len(table.rows)}, skipping")
        return
    if not placeholder_row_idxs or placeholder_row_idxs[0] >= len(table.rows):
        print(f"\u26a0 sec {sec_id}: placeholder out of range, skipping")
        return
    caption_tr = table.rows[caption_row_idx]._tr
    template_tr = copy.deepcopy(table.rows[placeholder_row_idxs[0]]._tr)

    # Remove placeholder rows (empty template photo rows) so only uploaded photos remain.
    # Do this after capturing `caption_tr` and `template_tr` references.
    for ri in sorted(set(placeholder_row_idxs), reverse=True):
        if 0 <= ri < len(table.rows):
            tr = table.rows[ri]._tr
            try:
                tr.getparent().remove(tr)
            except Exception:
                pass

    batches = [photos[i:i + n_cols] for i in range(0, len(photos), n_cols)]
    for batch in batches:
        new_tr = _make_tr_copy(template_tr)
        caption_tr.addprevious(new_tr)
        tcs = new_tr.findall(qn('w:tc'))
        for ci, img_path in enumerate(batch):
            if ci < len(tcs) and Path(img_path).exists():
                cell_obj = _Cell(tcs[ci], table)
                _insert_image_in_cell(
                    cell_obj,
                    str(img_path),
                    table=table,
                    col_idx=ci,
                    width_in=width_in,
                )


def _fill_thermal_photos(doc, photos, width_in=None):
    """Insert thermal camera photos into Table 25 cols 3 and 4 (0-indexed).
    Rows 1-4 are the data rows; cols 3 and 4 are the teplovizor image cells.
    Up to 8 photos: rows 1-4 x cols 3-4.
    """
    from docx.oxml.ns import qn
    from docx.table import _Cell

    if not photos or len(doc.tables) <= 25:
        return
    table = doc.tables[25]

    # Rows 1..4 are the 4 data rows (row 0 is header)
    photo_cells = []
    for r_idx in range(1, 5):          # rows 1,2,3,4
        if r_idx >= len(table.rows):
            break
        row = table.rows[r_idx]
        for c_idx in (3, 4):           # cols 3 and 4
            if c_idx >= len(row.cells):
                continue
            photo_cells.append((row.cells[c_idx], c_idx))

    for (cell, c_idx), img_path in zip(photo_cells, photos):
        if Path(img_path).exists():
            _insert_image_in_cell(
                cell,
                str(img_path),
                table=table,
                col_idx=c_idx,
                width_in=width_in,
            )


def _insert_all_photos(doc, photo_paths_by_sec, photo_widths):
    """Insert all section photos into the Word document."""
    thermal_photos = photo_paths_by_sec.get(10)  # sec_id 10 = thermal camera

    photo_sections = []
    for sec_id, photos in photo_paths_by_sec.items():
        if sec_id == 10 or sec_id not in PHOTO_SECTION_MAP:
            continue
        table_idx, _, caption_row_idx = PHOTO_SECTION_MAP[sec_id]
        width_in = photo_widths.get(sec_id, DEFAULT_PHOTO_WIDTHS.get(sec_id, 3.0))
        photo_sections.append((table_idx, caption_row_idx, sec_id, photos, width_in))

    # Process lower sections first so inserted rows do not shift later table indices.
    for _, _, sec_id, photos, width_in in sorted(
        photo_sections,
        key=lambda item: (item[0], item[1]),
        reverse=True,
    ):
        _fill_photo_table_section(doc, sec_id, photos, width_in=width_in)

    if thermal_photos:
        width_in = photo_widths.get(10, DEFAULT_PHOTO_WIDTHS.get(10, 3.0))
        _fill_thermal_photos(doc, thermal_photos, width_in=width_in)


def _insert_photos_by_placeholder(doc, photo_paths_by_sec, photo_widths=None):
    """Replace {{photo_secN_M}} or {{photo_N}} placeholders in template cells with photos.

    Primary format: {{photo_sec1_1}} → section 1, photo index 1 (1-based).
    Fallback format: {{photo_1}} → global sequential index across all sections.
    """
    import re

    if photo_widths is None:
        photo_widths = {}

    flat_photos = []
    for sec_id in sorted(photo_paths_by_sec.keys()):
        for path in photo_paths_by_sec[sec_id]:
            flat_photos.append((path, sec_id))

    pat_sec    = re.compile(r'^\{\{photo_sec(\d+)_(\d+)\}\}$')
    pat_global = re.compile(r'^\{\{photo_(\d+)\}\}$')

    for table in doc.tables:
        for row in table.rows:
            seen = set()
            col_idx = 0
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                txt = cell.text.strip()

                m = pat_sec.match(txt)
                if m:
                    sec_id = int(m.group(1))
                    n = int(m.group(2)) - 1  # 0-based
                    for para in cell.paragraphs:
                        para.clear()
                    sec_photos = photo_paths_by_sec.get(sec_id, [])
                    if n < len(sec_photos):
                        path = Path(sec_photos[n])
                        if path.exists():
                            _insert_image_in_cell(cell, str(path), table=table,
                                                  col_idx=col_idx, width_in=None)
                    col_idx += 1
                    continue

                m = pat_global.match(txt)
                if m:
                    n = int(m.group(1)) - 1  # 0-based
                    for para in cell.paragraphs:
                        para.clear()
                    if n < len(flat_photos):
                        path, sec_id = flat_photos[n]
                        path = Path(path)
                        if path.exists():
                            _insert_image_in_cell(cell, str(path), table=table,
                                                  col_idx=col_idx, width_in=None)
                    col_idx += 1
                    continue

                col_idx += 1


# ── Fill U-value compliance table (Table 22) ──────────────────────────────

def _clean_engineering_table(doc):
    """Table 4 (2.2 Muhandislik tizimlari): delete rows where both tavsif and izoh
    are empty after placeholder replacement; fill single-empty cell with '–'."""
    if 4 >= len(doc.tables):
        return
    table = doc.tables[4]
    rows_to_del = []
    for ri in range(1, len(table.rows)):   # skip header row 0
        cells = table.rows[ri].cells
        if len(cells) < 3:
            continue
        tavsif = cells[1].text.strip()
        izoh   = cells[2].text.strip()
        if not tavsif and not izoh:
            rows_to_del.append(ri)
        else:
            if not tavsif:
                _set_cell_text(cells[1], '–', font_size=9)
            if not izoh:
                _set_cell_text(cells[2], '–', font_size=9)
    for ri in reversed(rows_to_del):
        tr = table.rows[ri]._tr
        tr.getparent().remove(tr)


def _fill_u_table(doc, variables, form_data):
    """Overwrite the Moslik (compliance) column in Table 22."""
    if 22 >= len(doc.tables):
        return
    t = doc.tables[22]
    # Row 1=wall, 2=win, 3=roof, 4=floor
    keys = ['u_wall_status', 'u_win_status', 'u_roof_status', 'u_floor_status']
    defaults = ['Mos ✅', 'Mos ✅', 'Mos ✅', 'Mos ✅']
    for row_idx, (key, default) in enumerate(zip(keys, defaults), start=1):
        if row_idx >= len(t.rows):
            break
        cells = t.rows[row_idx].cells
        status = form_data.get(key, default)
        if len(cells) >= 7:
            _set_cell_text(cells[6], status, font_size=9)


# ── Insert charts after energy section ────────────────────────────────────

def _insert_charts(doc, chart_paths):
    """Insert chart images into Table 17 (8-rasm chart table).
    Table 17 structure: row 0 = chart cells, row 1 = more cells, row 2 = caption.
    3 pie charts go into row 0 cols 0,1,2; bar chart goes into row 1 col 0 (merged).
    """
    from docx.table import _Cell
    if 17 >= len(doc.tables):
        print("⚠ Chart table 17 not found")
        return
    table = doc.tables[17]
    # Collect unique cells from rows 0 and 1
    def unique_cells(row):
        seen, result = set(), []
        for cell in row.cells:
            if id(cell._tc) not in seen:
                seen.add(id(cell._tc))
                result.append(cell)
        return result

    slots = []
    for ri in range(min(2, len(table.rows))):
        slots.extend(unique_cells(table.rows[ri]))

    for cell, chart_path in zip(slots, chart_paths):
        w = _cell_actual_width_inches(cell)
        if w is None:
            w = _col_width_from_tblGrid(table, 0)
        if w is None:
            w = 3.0
        w = max(0.5, w - 0.16)
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run()
        try:
            run.add_picture(str(chart_path), width=Inches(w))
        except Exception as e:
            print(f"⚠ chart insert error: {e}")
            run.text = f"[Chart: {Path(chart_path).name}]"


# ── Measurement auto-generation helpers ───────────────────────────────────

def _rnd(val, spread, decimals=1, min_val=None):
    """Return val ± uniform(0, spread), rounded to decimals."""
    offset = random.uniform(-spread, spread)
    result = round(float(val or 0) + offset, decimals)
    if min_val is not None:
        result = max(min_val, result)
    return result


def _room_vars(data):
    """r1/r2 from form; r3/r4 auto-varied. All temp/hum formatted as 00,0."""
    v = {}
    for i in [1, 2]:
        t = float(data.get(f'r{i}_temp') or 22)
        h = float(data.get(f'r{i}_hum')  or 50)
        v[f'r{i}_temp'] = _fmt_meas(t)
        v[f'r{i}_hum']  = _fmt_meas(h)
        v[f'r{i}_lux']  = data.get(f'r{i}_lux', '')
    # Auto-generate r3 from r1, r4 from r2
    for src, dst in [(1, 3), (2, 4)]:
        t = float(data.get(f'r{src}_temp') or 22)
        h = float(data.get(f'r{src}_hum')  or 50)
        l = float(data.get(f'r{src}_lux')  or 300)
        v[f'r{dst}_temp'] = _fmt_meas(_rnd(t, 5, 1))
        v[f'r{dst}_hum']  = _fmt_meas(_rnd(h, 5, 1, min_val=10))
        v[f'r{dst}_lux']  = int(_rnd(l, 25, 0, min_val=50))
    return v


def _u_val(norm):
    """Return a U-value near the norm standard, varied ±0–0.35."""
    return round(norm + random.uniform(-0.35, 0.35), 2)


def _fmt_date(d):
    """Return date as DD.MM.YYYY for report display.
    Converts legacy YYYY-MM-DD input; passes DD.MM.YYYY through unchanged.
    """
    if not d:
        return ''
    if len(d) == 10 and d[4] == '-':          # YYYY-MM-DD
        y, m, day = d.split('-')
        return f"{day}.{m}.{y}"
    return d                                   # already DD.MM.YYYY


def _fmt_meas(val, decimals=1):
    """Format measurement value as 00,0 (comma decimal separator)."""
    try:
        return f"{float(val):.{decimals}f}".replace('.', ',')
    except Exception:
        return str(val) if val else '–'


# ── Build variables dict ───────────────────────────────────────────────────

def build_variables(data, energy, areas, case_no):
    fes_kw = int(data.get('fes_kw') or 10)
    fv = fes_values(fes_kw)

    geo_l = int(data.get('gelio_l') or 200)
    gv = gelio_values(geo_l)

    ariston_cnt = data.get('ariston_count', 1)
    ariston_kw  = data.get('ariston_kW', 2)
    ar = calc_ariston(ariston_cnt, ariston_kw)

    gas_2025_total = energy[2025]['summary']['total_gas']
    heat_area = float(data.get('heat_area') or areas['floor_area'] or 1)
    eff = calc_efficiency(gas_2025_total, heat_area)

    # BEE class A–G by signed % deviation from SHNQ norm 149 kWh/m²
    # (negative = uses less than norm = more efficient)
    _pct = eff['ec_pct_signed']
    if   _pct < -40:            ec_class = 'A'
    elif _pct <= -26:           ec_class = 'B'
    elif _pct <= -11:           ec_class = 'C'
    elif _pct <=   4:           ec_class = 'D'
    elif _pct <=  14:           ec_class = 'E'
    elif _pct <=  25:           ec_class = 'F'
    else:                       ec_class = 'G'

    # EC diff text
    diff_word = "kam" if eff['better'] else "ko'p"

    grid = data.get('grid', 'on-grid')
    # "batareyasi," only appears when FES type is hybrid
    batareyasi_text = "batareyasi," if 'hybrid' in grid.lower() else ''

    # fes_payb: show full payback formula instead of bare period
    _fes_mln   = fv['fes_mln']
    _fes_som   = fv['fes_som']
    _fes_payb  = fv['fes_payb']
    fes_payb_formula = (
        f"T = {_fes_mln} mln. soʻm ÷ {_fes_som:,.0f} soʻm = {_fes_payb}"
    )

    s23 = energy[2023]['summary']
    s24 = energy[2024]['summary']
    s25 = energy[2025]['summary']

    # Appliance totals
    apl_total_kwh = 0.0
    appl_kw_total = 0.0
    for i in range(1, 11):
        w   = float(data.get(f'apl{i}_w')   or 0)
        n   = float(data.get(f'apl{i}_n')   or 0)
        hrs = float(data.get(f'apl{i}_hrs') or 0)
        apl_total_kwh += w * n * hrs * 365 / 1000
        appl_kw_total += w * n / 1000

    # Door / window counts
    door_count = 0
    i = 1
    while data.get(f'door_n{i}') is not None:
        door_count += int(float(data.get(f'door_n{i}') or 0))
        i += 1

    win_count = 0
    i = 1
    while data.get(f'win_n{i}') is not None:
        win_count += int(float(data.get(f'win_n{i}') or 0))
        i += 1

    # FES panel area (5 m² per kW)
    fes_m2 = round(fes_kw * 5, 1)

    # Investment & savings totals
    tot_inv  = round(fv['fes_mln'] + gv['gelio_inv'], 1)
    tot_save = round((fv['fes_som'] + gv['gelio_som']) / 1e6, 2)

    # Ariston cost estimates (2025 tier-1 rate 600 som/kWh)
    _ar_rate = 600
    ar_kW_oy_c   = round(ar['month'] * _ar_rate)
    ar_kW_year_c = round(ar['year']  * _ar_rate)

    # CO2 total for 2025 (tonnes)
    co2_tot = round(s25['total_elec'] * 0.5 / 1000, 2)

    v = {
        # Case metadata
        'case_no':      case_no,
        'aud_date':     _fmt_date(data.get('aud_date', '')),
        'insp_date':    _fmt_date(data.get('insp_date', '')),
        'aud_name':     data.get('aud_name', ''),
        'aud_jshshr':   data.get('aud_jshshr', ''),

        # Location — mfy/street/house formatted for the template
        'region':       data.get('region', ''),
        'city':         data.get('city', ''),
        'mfy':          (f"{data.get('mfy', '')} MFY," if (data.get('mfy') or '').strip() else ''),
        'street':       (f"{data.get('street', '')} ko'chasi," if (data.get('street') or '').strip() else ''),
        'house':        (f"{data.get('house', '')} -uy" if (data.get('house') or '').strip() else ''),
        'lat':          data.get('lat', ''),
        'lon':          data.get('lon', ''),

        # Building
        'owner':        data.get('owner', ''),
        'residents':    data.get('residents', ''),
        'rooms':        data.get('rooms', ''),
        'floors':       data.get('floors', ''),
        'sections':     data.get('sections', ''),
        'area_total':   data.get('area_total') or areas['floor_area'],
        'yr_built':     data.get('yr_built', ''),
        'yr_renov':     data.get('yr_renov', ''),

        # Materials
        'wall_mat':     data.get('wall_mat', ''),
        'wall_thick':   data.get('wall_thick', ''),
        'wall_area':    areas['wall_net'],
        'roof_mat':     data.get('roof_mat', ''),
        'roof_area':    areas['roof_area'],
        'floor_mat':    data.get('floor_mat', ''),
        'floor_insul':  data.get('floor_insul', ''),
        'floor_area':   areas['floor_area'],
        'basement_mat': data.get('basement_mat') or "Yo'q",
        'basement_area':data.get('basement_area') or '—',

        # Windows & doors
        'win_area':     areas['win_area'],
        'door_area':    areas['door_area'],

        # Systems
        'heat_desc':    data.get('heat_desc', ''),
        'heat_note':    data.get('heat_note', ''),
        'hotw_desc':    data.get('hotw_desc', ''),
        'hotw_note':    data.get('hotw_note', ''),
        'light_desc':   data.get('light_desc', ''),
        'light_note':   data.get('light_note', ''),
        'appl_desc':    data.get('appl_desc', ''),
        'appl_note':    data.get('appl_note', ''),

        # Water heater
        'ariston_count':ariston_cnt,
        'ariston_kW':   ariston_kw,
        'ar_kW_day':    ar['day'],
        'ar_kW_oy':     ar['month'],
        'ar_kW_year':   ar['year'],

        # FES solar
        **fv,
        'grid':         grid,
        'fes_inv':      fv['fes_mln'],
        'fes_kwh':      fv['fes_kWh'],  # lowercase alias
        'fes_payb':     fes_payb_formula,
        'batareyasi,':  batareyasi_text,

        # Heliocollector
        **gv,
        'gelio_inv':    gv['gelio_inv'],
        'gelio_kwh':    gv['gelio_kwh'],

        # Energy efficiency
        'heat_area':    heat_area,
        'gas_2025':     int(gas_2025_total),
        'gas_95':       eff['gas_95'],
        'gas_heat':     eff['gas_heat'],
        'qov_fakt':     eff['qov_fakt'],
        'qov_to':       eff['qov_fakt'],    # alias
        'qov_fakt_val': eff['qov_fakt'],
        'ec_diff':      eff['ec_diff'],
        'ec_diff_percent': eff['ec_pct'],
        'e_from_mln':   round(s23['total_mln'], 2),
        'e_to_mln':     round(s25['total_mln'], 2),

        # Room measurements: r1/r2 from form, r3/r4 auto-varied ±5 / ±25
        **_room_vars(data),

        # Thermal U-values: auto-generated near standard ± random 0–0.35
        'u1_temp':      data.get('u1_temp', '21'),
        'u_wall':       _u_val(0.90),
        'u_win':        _u_val(2.80),
        'u_roof':       _u_val(0.80),
        'u_floor':      _u_val(1.00),

        # Materials (windows, doors, insulation)
        'win_mat':      data.get('win_mat', ''),
        'win_layers':   data.get('win_layers', ''),
        'wall_insul':   data.get('wall_insul', ''),
        'door_mat':     data.get('door_mat', ''),

        # Engineering systems (ventilation, water, cooling)
        'vent_desc':    data.get('vent_desc', ''),
        'vent_note':    data.get('vent_note', ''),
        'water_desc':   data.get('water_desc', ''),
        'water_note':   data.get('water_note', ''),
        'cool_desc':    data.get('cool_desc', ''),
        'cool_note':    data.get('cool_note', ''),
        'electric_reason': data.get('electric_reason', ''),

        # Account numbers
        'elec_account': data.get('elec_account', ''),
        'gas_account':  data.get('gas_account', ''),

        # Energy passport extras
        'light_kw':     data.get('light_kw', ''),
        'cool_kw':      data.get('cool_kw', ''),
        'gas_2023':     int(s23['total_gas']),
        'gas_2024':     int(s24['total_gas']),
        'e23_sum':      int(s23['total_elec']),
        'e24_sum':      int(s24['total_elec']),
        'e25_sum':      int(s25['total_elec']),
        '2025_kW':      int(s25['total_elec']),

        # Energy cost breakdowns
        'gk23_sum':     int(s23['total_gas_cost']),
        'gk24_sum':     int(s24['total_gas_cost']),
        'gk25_sum':     int(s25['total_gas_cost']),
        'tk23_sum':     int(s23['total_cost']),
        'tk24_sum':     int(s24['total_cost']),
        'tk25_sum':     int(s25['total_cost']),
        'e23_mln':      round(s23['total_elec_cost'] / 1e6, 3),
        'e24_mln':      round(s24['total_elec_cost'] / 1e6, 3),
        'e25_mln':      round(s25['total_elec_cost'] / 1e6, 3),

        # Appliances
        'apl_total':    round(apl_total_kwh, 1),
        'appl_kw':      round(appl_kw_total, 2),

        # Doors / windows counts
        'door_count':   door_count,
        'win_count':    win_count,

        # Efficiency & CO2
        'ec_class':     ec_class,
        'co2_tot':      co2_tot,

        # FES panel area
        'fes_m2':       fes_m2,

        # Investment totals
        'tot_inv':      tot_inv,
        'tot_save':     tot_save,

        # Ariston costs
        'ar_kW_oy_c':   ar_kW_oy_c,
        'ar_kW_year_c': ar_kW_year_c,
    }

    # Appliances (up to 10)
    for i in range(1, 11):
        for field in ['name', 'w', 'n', 'hrs']:
            v[f'apl{i}_{field}'] = data.get(f'apl{i}_{field}', '')

    return v


# ── Word report ────────────────────────────────────────────────────────────

def generate_word_report(variables, energy, chart_paths,
                         photo_paths_by_sec, photo_widths, out_path, form_data=None):
    src = BASE_DIR / 'template.docx'
    doc = Document(str(src))

    # Fill placeholders
    _replace_all(doc, variables)

    # Fill energy monthly tables — Table 10=2023, 12=2024, 14=2025, 16=comparison
    TABLE_IDX = {10: 2023, 12: 2024, 14: 2025}
    for t_idx, year in TABLE_IDX.items():
        if t_idx < len(doc.tables):
            _fill_energy_table(
                doc.tables[t_idx],
                energy[year]['rows'],
                energy[year]['summary'],
                year,
            )
    if 16 < len(doc.tables):
        _fill_comparison_table(doc.tables[16], energy)

    # Clean engineering systems table (2.2): delete empty rows, fill single-empty with "–"
    _clean_engineering_table(doc)

    # Fill U-value compliance column (Table 22)
    _fill_u_table(doc, variables, form_data or {})

    # Insert section photos via {{photo_N}} placeholders in template cells
    _insert_photos_by_placeholder(doc, photo_paths_by_sec, photo_widths)

    # Insert auto-generated charts
    if chart_paths:
        _insert_charts(doc, chart_paths)

    doc.save(str(out_path))
    return out_path


# ── Energy Passport ────────────────────────────────────────────────────────

def generate_passport(variables, out_path):
    src = BASE_DIR / 'passport_template.docx'
    doc = Document(str(src))
    _replace_all(doc, variables)

    # Fill passport table 4: annual energy (2023 and 2024)
    if 4 < len(doc.tables):
        t = doc.tables[4]
        data_rows = [
            (2023, variables.get('gas_2023', ''), variables.get('e23_sum', '')),
            (2024, variables.get('gas_2024', ''), variables.get('e24_sum', '')),
        ]
        for i, (year, gas, elec) in enumerate(data_rows):
            row_idx = i + 1
            if row_idx < len(t.rows):
                cells = t.rows[row_idx].cells
                if len(cells) >= 3:
                    _set_cell_text(cells[0], year)
                    _set_cell_text(cells[1], gas)
                    _set_cell_text(cells[2], elec)

    doc.save(str(out_path))
    return out_path


# ── Main entry point ───────────────────────────────────────────────────────

def _write_input_file(file_obj, dst):
    if hasattr(file_obj, 'save'):
        file_obj.save(str(dst))
        return
    if isinstance(file_obj, BinaryUpload):
        dst.write_bytes(file_obj.content)
        return
    raise TypeError(f"Unsupported photo input: {type(file_obj)!r}")


def _materialize_photo_inputs(photo_inputs, photos_dir):
    """Write uploaded or stored photos into a temp folder for report generation."""
    photo_paths_by_sec = {}
    stored_photos = {}

    print(f"_materialize_photo_inputs: {len(photo_inputs)} photo(s)")
    for (sec_id, n), file_obj in sorted(photo_inputs.items()):
        filename = getattr(file_obj, 'filename', '') or f"s{sec_id}_{n:02d}.jpg"
        ext = Path(filename).suffix or '.jpg'
        dst = photos_dir / f"s{sec_id}_{n:02d}{ext}"
        _write_input_file(file_obj, dst)
        raw_size = dst.stat().st_size if dst.exists() else 0
        print(f"  sec={sec_id} slot={n}: wrote {dst.name} ({raw_size} bytes)")
        dst = _convert_image_if_needed(dst)
        dst = _normalize_to_jpeg(dst)
        if not dst.exists() or dst.stat().st_size == 0:
            print(f"  ⚠ skipping sec={sec_id} slot={n}: file missing or empty after normalize")
            continue
        content = dst.read_bytes()
        stored_photos[(sec_id, n)] = BinaryUpload(
            filename=dst.name,
            content=content,
            mimetype=guess_mimetype(dst.name),
        )
        photo_paths_by_sec.setdefault(sec_id, []).append(dst)
        print(f"  ✓ sec={sec_id} slot={n}: ready as {dst.name} ({len(content)} bytes)")

    print(f"_materialize_photo_inputs: sections with photos: {sorted(photo_paths_by_sec.keys())}")
    return photo_paths_by_sec, stored_photos


def _generated_file_record(path, file_kind):
    path = Path(path)
    content = path.read_bytes()
    return {
        'filename': path.name,
        'file_kind': file_kind,
        'mime_type': guess_mimetype(path.name),
        'content': content,
    }


def generate_all(form_data, uploaded_photos, edit_case_name=None, created_by=None):
    """
    form_data: dict of all form fields
    uploaded_photos: dict {(sec_id, n): file-like object}
    edit_case_name: if set, update files in this existing case
    Returns: dict with case_name, case_no, files list
    """
    storage = get_storage()
    old_case_name = None

    owner_surname = (form_data.get('owner_surname') or 'CLIENT').upper().replace(' ', '_')

    if edit_case_name:
        existing_case = storage.get_case(edit_case_name)
        if not existing_case:
            raise ValueError(f"Case not found: {edit_case_name}")
        case_no = existing_case['case_no']
        case_name = f"EA-{case_no}_{owner_surname}"
        if case_name != edit_case_name:
            old_case_name = edit_case_name
        photo_inputs = storage.get_case_photos(edit_case_name)
    else:
        user_case_no = (form_data.get('case_number') or '').strip()
        case_no = user_case_no if user_case_no else storage.next_case_number()
        case_name = f"EA-{case_no}_{owner_surname}"
        photo_inputs = {}

    photo_inputs.update(uploaded_photos)

    photo_widths = {
        sec_id: float(form_data.get(f'photo_w{sec_id}') or DEFAULT_PHOTO_WIDTHS.get(sec_id, 3.0))
        for sec_id in range(1, 11)
    }

    serializable = {k: v for k, v in form_data.items()
                    if isinstance(v, (str, int, float, list, dict, type(None)))}

    # Calculate
    energy = calc_all(form_data)
    areas  = calc_areas(form_data)

    # Variables
    variables = build_variables(form_data, energy, areas, case_no)

    xlsx_name = f"{case_name}.xlsx"
    docx_name = f"{case_name}.docx"
    passport_name = f"{case_name}_passport.docx"

    with tempfile.TemporaryDirectory(prefix=f"{case_name}-") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        photos_dir = tmp_dir / 'photos'
        photos_dir.mkdir(parents=True, exist_ok=True)

        photo_paths_by_sec, stored_photos = _materialize_photo_inputs(photo_inputs, photos_dir)

        xlsx_path = tmp_dir / xlsx_name
        write_excel(form_data, energy, areas, xlsx_path)

        chart_paths = generate_charts(energy, tmp_dir)

        docx_path = tmp_dir / docx_name
        generate_word_report(
            variables,
            energy,
            chart_paths,
            photo_paths_by_sec,
            photo_widths,
            docx_path,
            form_data,
        )

        passport_path = tmp_dir / passport_name
        generate_passport(variables, passport_path)

        generated_files = [
            _generated_file_record(xlsx_path, 'excel'),
            _generated_file_record(docx_path, 'report'),
            _generated_file_record(passport_path, 'passport'),
        ]

    if old_case_name:
        storage.rename_case(old_case_name, case_name)

    storage.save_case(
        case_name=case_name,
        case_no=case_no,
        owner=form_data.get('owner', ''),
        form_data=serializable,
        files=generated_files,
        photos=stored_photos,
        created_by=created_by,
    )

    return {
        'success': True,
        'case_no': case_no,
        'case_name': case_name,
        'folder_name': case_name,
        'files': [xlsx_name, docx_name, passport_name],
    }
